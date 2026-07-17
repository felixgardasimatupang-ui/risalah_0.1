import json
import os
import sys
from datetime import datetime

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

AUDIO_EXTS = {
    ".mp3",
    ".mp4",
    ".m4a",
    ".wav",
    ".ogg",
    ".flac",
    ".aac",
    ".wma",
    ".mov",
    ".avi",
    ".mkv",
}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"}
DOC_EXTS = {".docx", ".doc", ".pdf", ".txt", ".rtf", ".md", ".csv", ".xlsx", ".xls", ".pptx"}
ALL_EXTS = AUDIO_EXTS | IMAGE_EXTS | DOC_EXTS


def classify(ext) -> None:
    if ext in AUDIO_EXTS:
        return "audio"
    if ext in IMAGE_EXTS:
        return "image"
    if ext in {".pdf", ".docx", ".doc", ".rtf"}:
        return "document"
    if ext in {".txt", ".md"}:
        return "text"
    if ext in {".csv", ".xlsx", ".xls"}:
        return "spreadsheet"
    return "other"


def scan_folder(folder_path) -> None:
    if not os.path.exists(folder_path):
        raise FileNotFoundError(f"Folder tidak ditemukan: {folder_path}")

    all_files = []
    for root, dirs, files in os.walk(folder_path):
        dirs.sort()
        for f in sorted(files):
            ext = os.path.splitext(f)[1].lower()
            if ext in ALL_EXTS:
                fp = os.path.join(root, f)
                all_files.append(
                    {
                        "path": fp,
                        "name": f,
                        "ext": ext,
                        "type": classify(ext),
                        "size_bytes": os.path.getsize(fp),
                    }
                )

    result = {
        "folder": os.path.abspath(folder_path),
        "scan_time": datetime.now().isoformat(),
        "total_files": len(all_files),
        "audio_files": [f for f in all_files if f["type"] == "audio"],
        "image_files": [f for f in all_files if f["type"] == "image"],
        "document_files": [f for f in all_files if f["type"] == "document"],
        "text_files": [f for f in all_files if f["type"] == "text"],
        "spreadsheet_files": [f for f in all_files if f["type"] == "spreadsheet"],
        "all_files": all_files,
    }

    print(
        f"Scan: {len(result['audio_files'])} audio, {len(result['image_files'])} gambar, "
        f"{len(result['document_files'])} dokumen, {len(result['text_files'])} teks, "
        f"{len(result['spreadsheet_files'])} spreadsheet = {result['total_files']} total"
    )

    scan_path = os.path.join(PROJECT_ROOT, "output", "scan_results.json")
    os.makedirs(os.path.dirname(scan_path), exist_ok=True)
    with open(scan_path, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
    return result


def read_txt(fp) -> None:
    with open(fp, encoding="utf-8", errors="replace") as f:
        return f.read()


def read_docx(fp) -> None:
    from docx import Document

    doc = Document(fp)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for t in doc.tables:
        for row in t.rows:
            line = " | ".join(c.text.strip() for c in row.cells)
            if line.strip():
                parts.append(line)
    return "\n".join(parts)


def read_pdf(fp) -> None:
    import fitz

    doc = fitz.open(fp)
    lines = []
    for page in doc:
        blocks = page.get_text("blocks")
        for b in blocks:
            if b[6] != 0:
                continue
            lines.append(b[4].strip())
    doc.close()
    if not "".join(lines).strip():
        import pdfplumber

        with pdfplumber.open(fp) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    lines.append(t)
    return "\n".join(lines)


def read_image(fp) -> None:
    from dotenv import load_dotenv

    load_dotenv()

    with open(fp, "rb") as f:
        import base64
        img_b64 = base64.b64encode(f.read()).decode()

    base_url = os.getenv("NINEROUTER_BASE_URL", "http://localhost:20128/v1")
    api_key = os.getenv("NINEROUTER_API_KEY", "")
    model = os.getenv("NINEROUTER_MODEL", "groq/llama-3.3-70b-versatile")

    import requests
    try:
        resp = requests.post(
            f"{base_url}/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": model,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": (
                                    "Ekstrak semua teks dari gambar ini dengan presisi tinggi. "
                                    "Jika ada tabel, format markdown. Bahasa Indonesia."
                                ),
                            },
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"},
                            },
                        ],
                    }
                ],
                "temperature": 0.1,
                "max_tokens": 2048,
                "stream": False,
            },
            timeout=60,
        )
        if resp.status_code == 200:
            data = resp.json()
            return data["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"  OCR via 9router gagal: {e}")

    print("  OCR via 9router gagal. Fallback ke pytesseract...")
    try:
        import pytesseract
        from PIL import Image

        img = Image.open(fp)
        text = pytesseract.image_to_string(img, lang="ind")
        return text.strip() if text.strip() else "[Kosong]"
    except ImportError:
        print("  pytesseract tidak terinstall. Install: pip install pytesseract")
        return "[OCR SKIP]"


def read_spreadsheet(fp) -> None:
    import csv

    ext = os.path.splitext(fp)[1].lower()
    lines = []
    if ext == ".csv":
        with open(fp, encoding="utf-8", errors="replace") as f:
            for row in csv.reader(f):
                lines.append(" | ".join(row))
    elif ext in (".xlsx", ".xls"):
        try:
            import openpyxl

            wb = openpyxl.load_workbook(fp, read_only=True, data_only=True)
            for name in wb.sheetnames:
                ws = wb[name]
                lines.append(f"\n--- Sheet: {name} ---")
                for row in ws.iter_rows(values_only=True):
                    r = " | ".join(str(c) if c is not None else "" for c in row)
                    if r.strip():
                        lines.append(r)
            wb.close()
        except ImportError:
            lines.append("[Install openpyxl: pip install openpyxl]")
        except Exception as e:
            lines.append(f"[Error xlsx: {e}]")
    return "\n".join(lines)


def extract_all_text(scan_result, output_dir=None) -> None:
    if output_dir is None:
        output_dir = os.path.join(PROJECT_ROOT, "output", "extracted_text")
    os.makedirs(output_dir, exist_ok=True)

    ctx = {
        "folder": scan_result["folder"],
        "extracted_at": datetime.now().isoformat(),
        "audio_sources": [],
        "image_sources": [],
        "document_sources": [],
        "all_text_combined": "",
    }

    for f in scan_result.get("audio_files", []):
        ctx["audio_sources"].append({"file": f["name"], "path": f["path"]})
    for f in scan_result.get("image_files", []):
        print(f"OCR: {f['name']}...")
        try:
            text = read_image(f["path"])
        except Exception as e:
            text = f"[Error OCR: {e}]"
        ctx["image_sources"].append({"file": f["name"], "text": text})
    for f in scan_result.get("document_files", []) + scan_result.get("text_files", []):
        print(f"Baca: {f['name']}...")
        try:
            ext = f["ext"]
            if ext == ".pdf":
                text = read_pdf(f["path"])
            elif ext in (".docx", ".doc"):
                text = read_docx(f["path"])
            elif ext in (".txt", ".md", ".rtf"):
                text = read_txt(f["path"])
            else:
                text = ""
        except Exception as e:
            text = f"[Error: {e}]"
        ctx["document_sources"].append({"file": f["name"], "text": text})
    for f in scan_result.get("spreadsheet_files", []):
        print(f"Sheet: {f['name']}...")
        try:
            text = read_spreadsheet(f["path"])
        except Exception as e:
            text = f"[Error sheet: {e}]"
        ctx["document_sources"].append({"file": f["name"], "text": text})

    parts = []
    if ctx["image_sources"]:
        parts.append("=== GAMBAR ===")
        for img in ctx["image_sources"]:
            parts.append(f"\n--- {img['file']} ---\n{img['text']}")
    if ctx["document_sources"]:
        parts.append("\n\n=== DOKUMEN ===")
        for d in ctx["document_sources"]:
            if d["text"].strip():
                parts.append(f"\n--- {d['file']} ---\n{d['text']}")
    ctx["all_text_combined"] = "\n".join(parts)

    json_path = os.path.join(output_dir, "extracted_text.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(ctx, f, indent=2, ensure_ascii=False)
    txt_path = os.path.join(output_dir, "extracted_text.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(ctx["all_text_combined"])
    print(f"Teks: {json_path}")
    return ctx


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python risalah/file_scanner.py <folder_path>")
        sys.exit(1)
    scan = scan_folder(sys.argv[1])
    if scan["total_files"] > 0:
        extract_all_text(scan)
