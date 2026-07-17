import json
import os
import sys
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

FONT_NAME = "Times New Roman"
FONT_SIZE_TITLE = 14
FONT_SIZE_HEADING = 12
FONT_SIZE_BODY = 12
FONT_SIZE_TABLE = 10

MARGIN_TOP = Cm(4)
MARGIN_LEFT = Cm(4)
MARGIN_RIGHT = Cm(3)
MARGIN_BOTTOM = Cm(3)

CLASSIFICATION_OPTIONS = ["BIASA", "TERBATAS", "RAHASIA"]


def set_page_setup(doc):
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


def set_run_font(run, size=FONT_SIZE_BODY, bold=False, italic=False, name=FONT_NAME):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    _rPr = run._r.get_or_add_rPr()  # noqa: N806
    _rFonts = OxmlElement("w:rFonts")  # noqa: N806
    _rFonts.set(qn("w:eastAsia"), name)
    _rPr.insert(0, _rFonts)


def set_paragraph_spacing(paragraph, line_spacing=1.5, space_after=None, space_before=None):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_after = space_after
    pf.space_before = space_before


def add_formal_paragraph(
    doc,
    text,
    size=FONT_SIZE_BODY,
    bold=False,
    italic=False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_after=None,
):
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_spacing(p, space_after=space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def set_table_borders(table):
    tbl = table._tbl
    _tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")  # noqa: N806
    borders = OxmlElement("w:tblBorders")
    for b in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{b}")
        for attr, val in [("val", "single"), ("sz", "4"), ("space", "0"), ("color", "000000")]:
            el.set(qn(f"w:{attr}"), val)
        borders.append(el)
    _tblPr.append(borders)


def set_cell_font(cell, text, size=FONT_SIZE_TABLE, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_spacing(p, line_spacing=1.0, space_after=Pt(2))
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_header_footer(doc, doc_number="_______________", classification="BIASA"):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(f"RISALAH RAPAT | {doc_number} | {classification}")
    set_run_font(run, size=8, italic=True)
    run.font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Halaman 1")
    set_run_font(run, size=8)


def add_title_block(doc, title="RISALAH RAPAT", classification="BIASA"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=1.0, space_after=Pt(2))
    run = p.add_run(classification)
    set_run_font(run, size=10, bold=True)
    run.font.color.rgb = RGBColor(200, 0, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=1.5, space_after=Pt(12))
    run = p.add_run(title)
    set_run_font(run, size=FONT_SIZE_TITLE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=1.0, space_after=Pt(18))
    run = p.add_run("—" * 30)
    set_run_font(run, size=8)


def add_metadata(doc, metadata):
    add_formal_paragraph(
        doc,
        "INFORMASI RAPAT",
        size=FONT_SIZE_HEADING,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=Pt(6),
    )

    fields = [
        ("Hari / Tanggal", metadata.get("tanggal", "_______________")),
        ("Waktu", metadata.get("waktu", "_______________")),
        ("Tempat", metadata.get("tempat", "_______________")),
        ("Acara", metadata.get("acara", "_______________")),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (label, value) in enumerate(fields):
        set_cell_font(table.rows[i].cells[0], label, size=FONT_SIZE_BODY, bold=True)
        set_cell_font(table.rows[i].cells[1], f"  : {value}", size=FONT_SIZE_BODY)
        table.rows[i].cells[0].width = Cm(4)
        table.rows[i].cells[1].width = Cm(12)


def add_attendees(doc, speakers):
    add_formal_paragraph(
        doc,
        "DAFTAR HADIR",
        size=FONT_SIZE_HEADING,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=Pt(6),
    )

    table = doc.add_table(rows=1, cols=5)
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["No", "Nama", "Jabatan", "Instansi", "Keterangan"]
    for i, h in enumerate(headers):
        set_cell_font(
            table.rows[0].cells[i],
            h,
            size=FONT_SIZE_BODY,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    seen = []
    for spk in speakers:
        name = spk.get("inferred_name", spk.get("label", ""))
        if name in seen or not name:
            continue
        seen.append(name)
        row = table.add_row().cells
        set_cell_font(row[0], str(len(seen)), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_font(row[1], name)
        set_cell_font(row[2], spk.get("inferred_role", ""))
        set_cell_font(row[3], "_______________")
        set_cell_font(row[4], "Hadir")

    if not seen:
        row = table.add_row().cells
        for i, v in enumerate(
            ["1", "_______________", "_______________", "_______________", "Hadir"]
        ):
            set_cell_font(row[i], v)


def add_transcript(doc, corrected, max_rows=200):
    add_formal_paragraph(
        doc,
        "ISI RISALAH",
        size=FONT_SIZE_HEADING,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=Pt(6),
    )

    table = doc.add_table(rows=1, cols=4)
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["No", "Waktu", "Pembicara", "Isi Pembicaraan"]
    for i, h in enumerate(headers):
        set_cell_font(
            table.rows[0].cells[i],
            h,
            size=FONT_SIZE_BODY,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    col_widths = [Cm(1), Cm(1.8), Cm(3.5), Cm(10.5)]
    for i, w in enumerate(col_widths):
        table.rows[0].cells[i].width = w

    for idx, seg in enumerate(corrected[:max_rows]):
        row = table.add_row().cells
        for i, w in enumerate(col_widths):
            row[i].width = w
        set_cell_font(row[0], str(idx + 1), size=FONT_SIZE_TABLE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_font(
            row[1], seg.get("time", ""), size=FONT_SIZE_TABLE, align=WD_ALIGN_PARAGRAPH.CENTER
        )
        set_cell_font(row[2], seg.get("speaker", ""), size=FONT_SIZE_TABLE)
        set_cell_font(row[3], seg.get("text", ""), size=FONT_SIZE_TABLE)

    if len(corrected) > max_rows:
        add_formal_paragraph(
            doc,
            f"Catatan: Transkrip lengkap ({len(corrected)} segmen) tersedia di file JSON output.",
            size=10,
            italic=True,
        )


def add_sections(doc, data):
    items = [
        ("POKOK BAHASAN", data.get("pokok_bahasan", [])),
        ("KEPUTUSAN RAPAT", data.get("keputusan_rapat", [])),
        ("KESIMPULAN", data.get("kesimpulan", [])),
        ("AGENDA RAPAT", data.get("agenda_rapat", [])),
    ]
    for title, list_items in items:
        if not list_items:
            continue
        add_formal_paragraph(
            doc,
            title,
            size=FONT_SIZE_HEADING,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            space_after=Pt(6),
        )
        for item in list_items:
            p = doc.add_paragraph(
                style="List Number" if title == "KEPUTUSAN RAPAT" else "List Bullet"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, line_spacing=1.5, space_after=Pt(3))
            run = p.add_run(item)
            set_run_font(run, size=FONT_SIZE_BODY)

    tl = data.get("tindak_lanjut", [])
    add_formal_paragraph(
        doc,
        "TINDAK LANJUT",
        size=FONT_SIZE_HEADING,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=Pt(6),
    )

    if tl:
        table = doc.add_table(rows=1, cols=4)
        set_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["No", "Tindakan", "PIC", "Batas Waktu"]):
            set_cell_font(
                table.rows[0].cells[i],
                h,
                size=FONT_SIZE_BODY,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
        for idx, item in enumerate(tl):
            row = table.add_row().cells
            set_cell_font(row[0], str(idx + 1), align=WD_ALIGN_PARAGRAPH.CENTER)
            if isinstance(item, dict):
                set_cell_font(row[1], item.get("tindakan", "-"))
                set_cell_font(row[2], item.get("pic", "-"))
                set_cell_font(row[3], item.get("batas_waktu", "-"))
            else:
                set_cell_font(row[1], str(item))
                set_cell_font(row[2], "-")
                set_cell_font(row[3], "-")
    else:
        add_formal_paragraph(doc, "- (tidak ada)", size=FONT_SIZE_BODY, italic=True)


def add_appendix(doc, data):
    dokumen_terkait = data.get("dokumen_terkait", [])
    dokumen_pendukung = data.get("dokumen_pendukung")
    dokumen_ringkasan = data.get("dokumen_ringkasan")

    if not dokumen_terkait and not dokumen_pendukung:
        return

    doc.add_page_break()
    add_formal_paragraph(
        doc,
        "LAMPIRAN",
        size=FONT_SIZE_TITLE,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(12),
    )

    if dokumen_terkait:
        add_formal_paragraph(
            doc,
            "Dokumen yang Dibahas dalam Rapat",
            size=FONT_SIZE_HEADING,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )
        for item in dokumen_terkait:
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p)
            run = p.add_run(item)
            set_run_font(run, size=FONT_SIZE_BODY)

    if dokumen_pendukung:
        add_formal_paragraph(
            doc,
            "Hasil Ekstraksi Dokumen Pendukung",
            size=FONT_SIZE_HEADING,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )

        if dokumen_ringkasan:
            r = dokumen_ringkasan.get("ringkasan", "")
            if r:
                add_formal_paragraph(doc, "Ringkasan:", bold=True, space_after=Pt(2))
                add_formal_paragraph(doc, r, space_after=Pt(6))

            for poin in dokumen_ringkasan.get("poin_penting", []):
                p = doc.add_paragraph(style="List Bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                set_paragraph_spacing(p)
                run = p.add_run(poin)
                set_run_font(run, size=FONT_SIZE_BODY)

            anggaran = dokumen_ringkasan.get("angka_anggaran", [])
            if anggaran:
                add_formal_paragraph(doc, "Anggaran:", bold=True, space_after=Pt(2))
                t = doc.add_table(rows=1, cols=2)
                set_table_borders(t)
                set_cell_font(t.rows[0].cells[0], "Item", bold=True)
                set_cell_font(t.rows[0].cells[1], "Jumlah", bold=True)
                for a in anggaran:
                    r = t.add_row().cells
                    set_cell_font(r[0], a.get("item", "-"))
                    set_cell_font(r[1], a.get("jumlah", "-"))


def add_signature(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p)
    now = datetime.now()
    bulan = [
        "Januari",
        "Februari",
        "Maret",
        "April",
        "Mei",
        "Juni",
        "Juli",
        "Agustus",
        "September",
        "Oktober",
        "November",
        "Desember",
    ]
    run = p.add_run(f"Jakarta, {now.day} {bulan[now.month - 1]} {now.year}")
    set_run_font(run, size=FONT_SIZE_BODY, italic=True)

    doc.add_paragraph()
    tt = doc.add_table(rows=1, cols=3)
    tt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, label in enumerate(["Pimpinan Rapat", "Mengetahui", "Sekretaris / Notulis"]):
        c = tt.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p)
        run = p.add_run(f"{label}\n\n\n\n(_______________)\n\nNama Jelas")
        set_run_font(run, size=FONT_SIZE_BODY)


def generate_risalah(
    enhanced,
    metadata=None,
    output_path=None,
    title="RISALAH RAPAT",
    doc_number="_______________",
    classification="BIASA",
):
    if classification not in CLASSIFICATION_OPTIONS:
        classification = "BIASA"

    if output_path is None:
        date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        od = os.path.join(PROJECT_ROOT, "output", "docs")
        os.makedirs(od, exist_ok=True)
        base = (metadata or {}).get("acara", "risalah").replace(" ", "_") if metadata else "risalah"
        output_path = os.path.join(od, f"{base}_{date_str}.docx")

    doc = Document()
    set_page_setup(doc)
    add_header_footer(doc, doc_number, classification)
    add_title_block(doc, title, classification)
    add_metadata(doc, metadata or {})
    doc.add_paragraph()
    add_attendees(doc, enhanced.get("speaker_identification", []))
    doc.add_paragraph()
    add_transcript(doc, enhanced.get("corrected_transcript", []))
    doc.add_paragraph()
    add_sections(doc, enhanced)
    add_appendix(doc, enhanced)
    add_signature(doc)

    doc.save(output_path)
    print(f"DOCX: {output_path}")
    return output_path


def generate_preview_text(enhanced, metadata=None, max_segments=50):
    lines = []
    lines.append("=" * 60)
    lines.append("PREVIEW RISALAH RAPAT")
    lines.append("=" * 60)
    lines.append("")

    m = metadata or {}
    lines.append(f"Acara   : {m.get('acara', '_______________')}")
    lines.append(f"Tanggal : {m.get('tanggal', '_______________')}")
    lines.append(f"Waktu   : {m.get('waktu', '_______________')}")
    lines.append(f"Tempat  : {m.get('tempat', '_______________')}")
    lines.append("")

    si = enhanced.get("speaker_identification", [])
    if si:
        lines.append("DAFTAR HADIR:")
        for s in si:
            lines.append(
                f"  - {s.get('inferred_name', s.get('label', ''))} ({s.get('inferred_role', '')})"
            )
        lines.append("")

    lines.append(
        "ISI RISALAH (preview {} dari {} segmen):".format(
            min(max_segments, len(enhanced.get("corrected_transcript", []))),
            len(enhanced.get("corrected_transcript", [])),
        )
    )
    lines.append("-" * 60)
    for idx, seg in enumerate(enhanced.get("corrected_transcript", [])[:max_segments]):
        lines.append(
            f"[{seg.get('time', '00:00')}] {seg.get('speaker', '')}: {seg.get('text', '')}"
        )
        if idx < max_segments - 1:
            lines.append("")
    if len(enhanced.get("corrected_transcript", [])) > max_segments:
        lines.append(
            f"\n... dan {len(enhanced['corrected_transcript']) - max_segments} segmen lainnya."
        )

    for section, key in [
        ("POKOK BAHASAN", "pokok_bahasan"),
        ("KEPUTUSAN", "keputusan_rapat"),
        ("KESIMPULAN", "kesimpulan"),
        ("TINDAK LANJUT", "tindak_lanjut"),
    ]:
        items = enhanced.get(key, [])
        if items:
            lines.append(f"\n{section}:")
            for item in items:
                if isinstance(item, dict):
                    lines.append(
                        f"  - {item.get('tindakan', '')} | PIC: {item.get('pic', '')} | {item.get('batas_waktu', '')}"
                    )
                else:
                    lines.append(f"  - {item}")

    lines.append("")
    lines.append("=" * 60)
    lines.append("Preview ini adalah draft. Jalankan tanpa --preview untuk DOCX final.")
    lines.append("=" * 60)

    return "\n".join(lines)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python risalah/docx_generator.py <enhanced_json> [metadata_json] [title]")
        print("  --preview : tampilkan preview TXT saja")
        sys.exit(1)

    preview = "--preview" in sys.argv
    args = [a for a in sys.argv[1:] if a != "--preview"]

    with open(args[0]) as f:
        data = json.load(f)
    meta = None
    if len(args) > 1:
        with open(args[1]) as f:
            meta = json.load(f)
    title = args[2] if len(args) > 2 else "RISALAH RAPAT"

    if preview:
        print(generate_preview_text(data, meta))
    else:
        generate_risalah(data, meta, title=title)
